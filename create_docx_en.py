import sys
import subprocess

def install_and_create():
    try:
        import docx
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
    
    doc = docx.Document()
    doc.add_heading('Additional Design Details for AI Rehab Brace Project', 0)
    
    doc.add_heading('1. Key Features & Functionalities', level=1)
    doc.add_paragraph('- Real-time muscle signal (sEMG) and movement (IMU) monitoring.\n- AI-driven muscle fatigue detection and adaptive resistance/assistance.\n- Contextual AI (Camera Vision) for preemptive grasp posture recommendations.\n- Mobile Application for patients to track progress, effort scores, and daily routines.\n- Tele-rehab Dashboard for physiotherapists to monitor remote patients.')

    doc.add_heading('2. Success Metrics', level=1)
    doc.add_paragraph('- AI fatigue detection accuracy: > 90%.\n- System response latency: < 100ms.\n- Estimated production cost: Under $150.\n- High patient adherence rate for at-home rehabilitation exercises.')

    doc.add_heading('3. Cost Estimation (POC Prototype)', level=1)
    doc.add_paragraph('- sEMG & IMU Sensors: ~$50\n- Microcontroller (ESP32/Raspberry Pi): ~$15\n- Mini Camera: ~$15\n- Actuators/Motors: ~$30\n- 3D Printed Casing & Materials: ~$20\n=> Total Estimated Hardware Cost: ~$130 (Highly affordable compared to thousands of dollars in the market).')

    doc.add_heading('4. Convenient / Value', level=1)
    doc.add_paragraph('- Convenience: Patients can perform rehabilitation safely at home without daily hospital visits.\n- Value: The AI acts as a "personal trainer," preventing injuries caused by incorrect posture or overexertion. Physiotherapists receive real-time data to adjust treatment plans remotely.')

    doc.add_heading('5. Related Content & Differentiation', level=1)
    doc.add_paragraph('- 2024 (AI in diagnosis): Many projects focused on diagnosis. Our project goes further into "intervention and therapeutics".\n- 2022 (Support movement): Existing movement support devices (exoskeletons) are often bulky and expensive. The differentiator for this project is Edge AI (Intel OpenVINO), which makes the device compact, affordable, and provides instant feedback.')

    doc.save('Thiet_Ke_Bo_Sung_EN.docx')

if __name__ == "__main__":
    install_and_create()
