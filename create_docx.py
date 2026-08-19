import sys
import subprocess

def install_and_create():
    try:
        import docx
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
    
    doc = docx.Document()
    doc.add_heading('Bổ sung các phần còn thiếu cho Dự án AI Rehab Brace', 0)
    
    doc.add_heading('1. Key Features & Functionalities (Các tính năng chính)', level=1)
    doc.add_paragraph('- Giám sát tín hiệu cơ (sEMG) và chuyển động (IMU) theo thời gian thực.\n- AI phát hiện mức độ mỏi cơ và tự động điều chỉnh lực cản/trợ lực (Adaptive Resistance).\n- AI Nhận diện ngữ cảnh (Camera Vision) tự động đề xuất tư thế nắm đồ vật.\n- Ứng dụng di động (Mobile App) cho bệnh nhân theo dõi tiến độ và điểm nỗ lực.\n- Bảng điều khiển từ xa (Tele-rehab Dashboard) cho bác sĩ vật lý trị liệu.')

    doc.add_heading('2. Success Metrics (Thước đo thành công)', level=1)
    doc.add_paragraph('- Độ chính xác của mô hình AI nhận diện mỏi cơ: > 90%.\n- Độ trễ phản hồi của hệ thống: < 100ms.\n- Chi phí sản xuất dự kiến: Dưới 150 USD.\n- Tỷ lệ tuân thủ bài tập của bệnh nhân khi sử dụng thiết bị tại nhà.')

    doc.add_heading('3. Cost Estimation (Ước tính chi phí - Bản mẫu POC)', level=1)
    doc.add_paragraph('- Cảm biến sEMG & IMU: ~50 USD\n- Vi điều khiển (ESP32/Raspberry Pi): ~15 USD\n- Camera mini: ~15 USD\n- Động cơ trợ lực (Actuators): ~30 USD\n- Vỏ 3D in và vật liệu khác: ~20 USD\n=> Tổng chi phí phần cứng ước tính: ~130 USD (Rất rẻ so với thị trường hàng ngàn đô).')

    doc.add_heading('4. Convenient / Value (Sự tiện lợi và Giá trị mang lại)', level=1)
    doc.add_paragraph('- Tiện lợi: Bệnh nhân có thể tập phục hồi chức năng tại nhà một cách an toàn mà không cần đến bệnh viện mỗi ngày.\n- Giá trị: AI đóng vai trò như một "huấn luyện viên cá nhân", ngăn ngừa chấn thương do tập sai tư thế hoặc quá sức. Bác sĩ có dữ liệu thực tế để điều chỉnh phác đồ điều trị từ xa.')

    doc.add_heading('5. Related Content (Các dự án liên quan & Sự khác biệt)', level=1)
    doc.add_paragraph('- Năm 2024 (AI in diagnosis): Nhiều dự án tập trung vào chẩn đoán. Dự án của chúng ta đi xa hơn vào "can thiệp và hỗ trợ điều trị" (Therapeutics).\n- Năm 2022 (Support movement): Các thiết bị hỗ trợ di chuyển (exoskeleton) thường cồng kềnh và đắt đỏ. Điểm khác biệt của dự án này là tính toán AI tại biên (Edge AI - Intel OpenVINO) giúp thiết bị nhỏ gọn, giá rẻ và phản hồi tức thì.')

    doc.save('Thiet_Ke_Bo_Sung.docx')

if __name__ == "__main__":
    install_and_create()
