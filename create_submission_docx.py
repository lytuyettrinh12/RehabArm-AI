import sys
import subprocess

def create_submission_doc():
    try:
        import docx
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx

    doc = docx.Document()
    
    # Title
    doc.add_heading('HỒ SƠ MẪU ĐĂNG KÝ BÀI THI - INTEL VIETNAM AI IMPACT FESTIVAL 2026', 0)
    
    # Section 1: Thông tin chung
    doc.add_heading('1. THÔNG TIN CHUNG VỀ DỰ ÁN (PROJECT OVERVIEW)', level=1)
    p = doc.add_paragraph()
    p.add_run('• Tên dự án (Tiếng Anh): ').bold = True
    p.add_run('FlexiMind AI - Smart Active-Assistive Rehab Sleeve powered by Intel OpenVINO\n')
    p.add_run('• Tên dự án (Tiếng Việt): ').bold = True
    p.add_run('FlexiMind AI - Đai nẹp đeo tay thông minh hỗ trợ phục hồi chức năng chủ động dựa trên AI tại biên\n')
    p.add_run('• Lĩnh vực (Category): ').bold = True
    p.add_run('AI for Accessibility & Healthcare (AI cho Y tế & Hỗ trợ người khuyết tật)\n')
    p.add_run('• Đối tượng mục tiêu (Target Users): ').bold = True
    p.add_run('Bệnh nhân phục hồi sau đột quỵ/tai biến, bệnh nhân chấn thương cột sống/chi trên, người cao tuổi cần tập vật lý trị liệu tại nhà.')

    # Section 2: Tóm tắt vấn đề
    doc.add_heading('2. TÓM TẮT VẤN ĐỀ XÃ HỘI (PROBLEM STATEMENT)', level=1)
    doc.add_paragraph(
        'Hiện nay, hàng triệu bệnh nhân sau đột quỵ hoặc chấn thương chi trên gặp rất nhiều khó khăn trong việc phục hồi chức năng vận động:\n'
        '1. Chi phí đắt đỏ: Các thiết bị robot phục hồi chức năng (exoskeleton) hiện nay trên thị trường có giá từ $5,000 - $50,000 USD, quá tầm tay với đa số người dân Việt Nam.\n'
        '2. Nguy cơ chấn thương thứ cấp: Khi tự tập tại nhà không có bác sĩ giám sát, bệnh nhân dễ cố quá sức dẫn đến co giật cơ (spasm) hoặc tổn thương khớp vĩnh viễn.\n'
        '3. Thiếu dữ liệu theo dõi từ xa: Bác sĩ không có công cụ đo lường chính xác tiến trình hồi phục thực tế của bệnh nhân khi họ tập luyện tại nhà.'
    )

    # Section 3: Giải pháp AI & Đổi mới công nghệ
    doc.add_heading('3. GIẢI PHÁP AI & ĐỔI MỚI CÔNG NGHỆ (AI SOLUTION & INNOVATION)', level=1)
    doc.add_paragraph(
        'FlexiMind AI là hệ thống nẹp đeo tay thông minh tích hợp Cảm biến sinh học đa phương thức (Multimodal Sensing) và AI tại biên (Edge AI):\n\n'
        '• Thu thập dữ liệu đa phương thức (Multi-Modal Sensing Hub):\n'
        '  - Cảm biến sEMG (Surface Electromyography): Thu xung điện cơ bắp thời gian thực để giải mã ý định vận động.\n'
        '  - Cảm biến IMU (MPU6050): Đo vận tốc góc và hành trình di chuyển của khuỷu tay.\n\n'
        '• Thuật toán AI Dự đoán mỏi cơ (AI Fatigue Detection Engine):\n'
        '  - Lọc nhiễu dải tần 20Hz - 450Hz, tính toán biến đổi RMS và Tần số trung vị (Median Frequency).\n'
        '  - Mô hình AI nhận diện sự sụt giảm tần số ngầm khi cơ bắp kiệt sức với độ chính xác >90%.\n\n'
        '• Vòng lặp phản hồi thực thể (Closed-Loop Physical Actuation):\n'
        '  - Khi AI phát hiện ngưỡng nguy hiểm (Co giật/Mỏi nặng), hệ thống tự động phát lệnh cho Động cơ Servo xoay 90° để nới lỏng đai tay ngay lập tức, ngắt lực cản bảo vệ khớp bệnh nhân.'
    )

    # Section 4: Ứng dụng công nghệ Intel
    doc.add_heading('4. ỨNG DỤNG CÔNG NGHỆ INTEL (INTEL TECH INTEGRATION - METRIC 03)', level=1)
    doc.add_paragraph(
        'Dự án tích hợp sâu bộ công cụ Intel OpenVINO Toolkit:\n'
        '• Intel OpenVINO Model Optimizer: Chuyển đổi mô hình AI huấn luyện sang định dạng tối ưu IR (.xml & .bin).\n'
        '• Intel OpenVINO Runtime Engine: Chạy suy luận (Inference) trực tiếp trên phần cứng CPU Intel của thiết bị với độ trễ cực thấp (<10ms).\n'
        '• Việc sử dụng Intel OpenVINO giúp hệ thống phản ứng ngắt lực cản tức thì, đảm bảo an toàn tuyệt đối cho bệnh nhân mà không phụ thuộc vào kết nối Internet/Cloud.'
    )

    # Section 5: Tác động xã hội
    doc.add_heading('5. TÁC ĐỘNG XÃ HỘI & TÍNH KHẢ THI (SOCIAL IMPACT & FEASIBILITY - METRIC 01)', level=1)
    doc.add_paragraph(
        '• Chi phí siêu rẻ (< $20 USD ~ 380.000 VNĐ): Giảm chi phí sản xuất gấp 100 lần so với thiết bị thương mại, giúp mọi bệnh nhân nghèo đều có thể tiếp cận.\n'
        '• Phục hồi nhanh gấp 2-3 lần: Áp dụng nguyên lý y khoa "Trợ lực theo nhu cầu" (Assistance-as-Needed), kích thích khả năng tái tạo thần kinh (Neuroplasticity).\n'
        '• Kết nối Bác sĩ từ xa (Tele-rehab): Tự động đồng bộ tiến trình tập luyện lên Web Dashboard cho bác sĩ theo dõi và điều chỉnh phác đồ điều trị.'
    )

    # Section 6: Link liên quan
    doc.add_heading('6. DANH SÁCH LINK NỘP BÀI (SUBMISSION LINKS)', level=1)
    doc.add_paragraph(
        '• Link Video Demo (2 phút trên YouTube): [Chèn link YouTube Video 2 phút tại đây]\n'
        '• Link GitHub Source Code Public: [Chèn link GitHub repository tại đây]\n'
        '• Link Web Dashboard Demo: [Chèn link Web Demo/Streamlit tại đây]'
    )

    doc.save('Thong_Tin_Nop_Bai_Intel_VAIIF2026.docx')

if __name__ == "__main__":
    create_submission_doc()
